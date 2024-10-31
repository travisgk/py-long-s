"""
Filename: _convert.py
Description: This file contains the primary functions 
             of converting given text.

Author: TravisGK
Version: 1.0.1a

License: MIT License
"""

try:
    import docx
    _docx_available = True
except ImportError:
    _docx_available = False

try:
    import odf.opendocument
    import odf.text
    _odf_available = True
except ImportError:
    _odf_available = False

import multiprocessing
import os
import re
from functools import partial
from ._split_words import *
from ._simple_conversions import *
from ._german_conversion import enable_debug_text, convert_german_word
from ._german_lists import (
    get_german_dicts,
    using_developer_mode,
    enable_developer_mode,
    sort_words,
    load_dicts,
)


def get_conversion_func(lang: str):
    """Returns the conversion function used for a particular language."""
    if lang == "en":
        return convert_english_word
    elif lang == "fr":
        return convert_french_word
    elif lang == "de":
        load_dicts()
        return convert_german_word
    elif lang == "es":
        return convert_spanish_word
    elif lang == "it":
        return convert_italian_word
    return None


def convert_text_file(src_path: str, dst_path: str = None, lang: str = "en"):
    """
    Loads a text file, converts all the text to use the long S,
    then saves it to an output text file.

    Parameters:
    src_path (string): path of the source text file.
    dst_path (string): path of the destination text file.
    lang (string): the language code for the text.
    """
    if not os.path.isfile(src_path):
        print(f"Could not find .txt file \"{src_path}\"")
        return

    with open(src_path, "r", encoding="utf-8") as input_file:
        content = input_file.read()
        results = convert(content, lang)

    if dst_path is None:
        dst_path = src_path[:-4] + "-long-s.txt"
    with open(dst_path, "w", encoding="utf-8") as output_file:
        output_file.write(results)


def convert_docx_file(src_path: str, dst_path: str = None, lang: str = "en"):
    """
    Loads a .docx file, converts all the text to use the long S,
    then saves it to an output .docx file.

    Parameters:
    src_path (string): path of the source .docx file.
    dst_path (string): path of the destination .docx file.
    lang (string): the language code for the text.
    """
    if not _docx_available:
        print(
            "To convert .docx files, the docx library is needed." \
            "Please run:\n\tpip install python-docx\n"
        )
        return

    _, extension = os.path.splitext(src_path)
    if not os.path.isfile(src_path):
        print(f"Could not find {extension} file \"{src_path}\"")
        return

    # opens the .docx file.
    doc = docx.Document(src_path)

    # gets all the displayed text contents.
    flat_runs = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            flat_runs.append(run.text)

    # puts all the text into one string
    # so the program can use multiprocessing,
    # then converts the string.
    SEPARATOR = "\uE030"
    whole_str = SEPARATOR.join(flat_runs)
    convert_str = convert(whole_str, lang=lang)

    # breaks the one giant string back into a list,
    # then lays each element back into the document
    # in order to replace text with the conversions.
    flat_runs = convert_str.split(SEPARATOR)
    current_run_i = 0
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = flat_runs[current_run_i]
            current_run_i += 1

    # saves the .docx.
    if dst_path is None:
        dst_path = src_path[:-len(extension)] + "-long-s" + extension
    doc.save(dst_path)


def convert_odf_file(src_path: str, dst_path: str = None, lang: str = "en"):
    """
    Loads an .odf file, converts all the text to use the long S,
    then saves it to an output .odf file.

    Parameters:
    src_path (string): path of the source .odf file.
    dst_path (string): path of the destination .odf file.
    lang (string): the language code for the text.
    """
    if not _odf_available:
        print(
            "To convert .odf files, the odf library is needed." \
            "Please run:\n\tpip install odfpy\n"
        )
        return

    _, extension = os.path.splitext(src_path)
    if not os.path.isfile(src_path):
        print(f"Could not find {extension} file \"{src_path}\"")
        return

    # opens the .odf file.
    doc = odf.opendocument.load(src_path) 
    text_elements = doc.getElementsByType(odf.text.P)

    # gets all the displayed text contents.
    flat_runs = []
    for paragraph in text_elements:
        _process_text_nodes(paragraph, flat_runs, replace_text=False)

    # puts all the text into one string
    # so the program can use multiprocessing,
    # then converts the string.
    SEPARATOR = "\uE030"
    whole_str = SEPARATOR.join(flat_runs)
    convert_str = convert(whole_str, lang=lang)

    # breaks the one giant string back into a list,
    # then lays each element back into the document
    # in order to replace text with the conversions.
    flat_runs = convert_str.split(SEPARATOR)
    for paragraph in text_elements:
        _process_text_nodes(paragraph, flat_runs, replace_text=True)

    # saves the .odf.
    if dst_path is None:
        dst_path = src_path[:-len(extension)] + "-long-s" + extension
    doc.save(dst_path)


# Helper function to gather and replace text recursively in spans.
def _process_text_nodes(node, text_list, replace_text=False):
    for child in node.childNodes:
        if child.nodeType == child.TEXT_NODE:  # Plain text
            if not replace_text:
                text_list.append(child.data)
            elif len(text_list) > 0:
                child.data = text_list.pop(0)
        elif child.nodeType == child.ELEMENT_NODE:
            # recursively processes child elements.
            _process_text_nodes(child, text_list, replace_text)


def convert(text: str, lang: str = "en"):
    """
    Places the long s (ſ) in a sentence and returns it.

    Parameters:
    text (str): the string to convert into archaeic spelling.
    lang (str): the language code for <text>. "en", "es", "fr", "it", or "de".

    Returns:
    str: text with the long s (ſ) placed.
    """
    CHUNK_SIZE = 30
    MULTIPROCESS_THRESHOLD = 50  # wordcount that triggers multiprocessing.

    convert_func = get_conversion_func(lang)

    if convert_func is None:
        print(f'language "{lang}" not found.' "the options are: en, fr, de, es, it.")
        return text

    words_with_indices = split_words_with_indices(text, lang)

    if len(words_with_indices) > MULTIPROCESS_THRESHOLD:
        # uses multiprocessing to convert massive amounts of words.
        non_words = {}
        next_new_index = 0
        for old_word, start_index in words_with_indices:
            if next_new_index != start_index:
                non_words[next_new_index] = text[next_new_index:start_index]
                next_new_index = start_index
            next_new_index = start_index + len(old_word)

        if next_new_index < len(text):
            non_words[next_new_index] = text[next_new_index:]

        # breaks all the words to be converted into dictionary chunks.
        chunks = []
        for i in range(0, len(words_with_indices), CHUNK_SIZE):
            chunks.append({})
            for j in range(i, min(len(words_with_indices), i + CHUNK_SIZE)):
                old_word, start_index = words_with_indices[j]
                chunks[-1][start_index] = old_word
        last_used_index = (len(words_with_indices) // CHUNK_SIZE) * CHUNK_SIZE

        # converts every chunk using multiprocessing.
        if lang == "de":
            dicts = get_german_dicts()
            process_func = partial(
                _process_german_chunk, convert_func=convert_func, dicts=dicts
            )
        else:
            process_func = partial(_process_chunk, convert_func=convert_func)

        with multiprocessing.Pool() as pool:
            results = pool.map(process_func, chunks)

        # combines all the results.
        combined_result = {k: v for r in results for k, v in r.items()}
        for key, value in non_words.items():
            combined_result[key] = value

        mapped_list = sorted(combined_result.items())
        text = "".join(substr for _, substr in mapped_list)

        return text

    else:
        # uses one thread to directly convert some words.
        # converts each word individually.
        words = split_words_with_indices(text, lang)
        if lang == "de":
            dicts = get_german_dicts()
            for old_word, start_index in words:
                new_word = convert_func(old_word, dicts)

                if old_word != new_word:
                    text = (
                        text[:start_index]
                        + new_word
                        + text[start_index + len(new_word) :]
                    )
        else:
            for old_word, start_index in words:
                new_word = convert_func(old_word)

                if old_word != new_word:
                    text = (
                        text[:start_index]
                        + new_word
                        + text[start_index + len(new_word) :]
                    )

        return text


def _process_german_chunk(chunk, convert_func, dicts):
    """
    Returns a dict of every word contained in
    the given German chunk after being converted.
    """
    return {key: convert_func(value, dicts=dicts) for key, value in chunk.items()}


def _process_chunk(chunk, convert_func):
    """
    Returns a dict of every word contained in the given chunk after being converted.
    """
    return {key: convert_func(value) for key, value in chunk.items()}